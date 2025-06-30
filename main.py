from fastapi import FastAPI, HTTPException, UploadFile, File
from pydantic import BaseModel
import requests
import os
import logging
import time
import re
from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
import chromadb
from langchain_community.document_loaders import TextLoader, CSVLoader
from langchain_core.documents import Document
import pdfplumber
from docx import Document as DocxDocument


logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


load_dotenv()

class QueryRequest(BaseModel):
    question: str

app = FastAPI()

@app.get("/")
async def root():
    return {"message": "RAG API is running"}


collection_name = "assistant_collection"
try:
    client = chromadb.PersistentClient(path="/app/chroma")
    collection = client.get_or_create_collection(collection_name)
except Exception as e:
    logger.error(f"Failed to initialize Chroma DB: {e}")
    raise


if not collection.get()['ids']:
    collection.delete(ids=collection.get()['ids']) if collection.get()['ids'] else None
    logger.info("Cleared existing collection data")


try:
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    logger.info("Embeddings model loaded successfully")
except Exception as e:
    logger.error(f"Failed to load embeddings model: {e}")
    class DummyEmbeddings:
        def embed_documents(self, texts):
            return [[0.0] * 384 for _ in texts]
        def embed_query(self, text):
            return [0.0] * 384
    embeddings = DummyEmbeddings()
    logger.info("Using dummy embeddings as fallback")


def generate_questions(text_chunk):
    API_KEY = os.getenv("GOOGLE_API_KEY")
    if not API_KEY:
        logger.error("GOOGLE_API_KEY not found")
        return []
    BASE_URL = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={API_KEY}"
    text_chunk = text_chunk[:2000].strip()
    payload = {"contents": [{"parts": [{"text": f"Сгенерируйте ровно 3 кратких вопроса на русском языке для этого юридического текста. Вопросы должны быть пронумерованы (1., 2., 3.) и содержать только текст вопроса без пояснений: {text_chunk}"}]}]}
    max_retries = 20
    retry_count = 0
    wait_time = 30

    while retry_count < max_retries:
        try:
            response = requests.post(BASE_URL, headers={"Content-Type": "application/json"}, json=payload, timeout=60)
            response.raise_for_status()
            questions_text = response.json()["candidates"][0]["content"]["parts"][0]["text"]
            questions = []
            lines = [line.strip() for line in questions_text.split("\n") if line.strip()]
            for line in lines:
                if line.startswith("1.") or line.startswith("2.") or line.startswith("3."):
                    questions.append(line)
                if len(questions) == 3:
                    break
            return questions if questions else []
        except requests.exceptions.RequestException as e:
            logger.error(f"Ошибка генерации вопросов (попытка {retry_count + 1}/{max_retries}): {e}")
            # Расширить обработку ошибок
            if any(error in str(e).lower() for error in ["429", "timeout", "remote disconnected", "connection aborted"]):
                logger.info(f"Проблема с подключением или лимит API. Ожидание {wait_time} секунд перед повторной попыткой...")
                time.sleep(wait_time)
                retry_count += 1
            else:
                break
    logger.error("Достигнуто максимальное число попыток. Возвращено пустой список вопросов.")
    return []


def clean_pdf_text(file_path):
    try:
        text = ""
        page_numbers = []
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                page_text = page.extract_text() or ""
                if not page_text.strip():
                    continue
                lines = [line.strip() for line in page_text.split('\n') if line.strip()]
                cleaned_text = ' '.join(lines)
                if len(cleaned_text.strip()) < 100:
                    continue
                text += f"\n[Page {page_num}]\n{cleaned_text}\n"
                page_numbers.append(str(page_num))
        if not text.strip():
            return "", {"source": os.path.basename(file_path)}
        metadata = {"source": os.path.basename(file_path), "page_numbers": ",".join(page_numbers)}
        logger.info(f"Извлечено текста из {file_path}: {len(text)} символов")
        return text, metadata
    except Exception as e:
        logger.error(f"Ошибка обработки PDF {file_path}: {e}")
        return "", {"source": os.path.basename(file_path)}


def load_docx_text(file_path):
    try:
        doc = DocxDocument(file_path)
        text = "\n".join([para.text.strip() for para in doc.paragraphs if para.text.strip()])
        if not text:
            logger.warning(f"Нет пригодного текста в {os.path.basename(file_path)}")
            return "", {"source": os.path.basename(file_path)}
        logger.info(f"Извлечено {len(text)} символов из {os.path.basename(file_path)}")
        return text, {"source": os.path.basename(file_path)}
    except Exception as e:
        logger.error(f"Ошибка загрузки DOCX {file_path}: {e}")
        return "", {"source": os.path.basename(file_path)}


def load_documents(docs_dir):
    logger.info(f"Проверка папки с документами: {os.path.abspath(docs_dir)}")
    documents = []
    loaders = {
        ".txt": TextLoader,
        ".pdf": lambda x: clean_pdf_text(x),
        ".docx": load_docx_text,
        ".csv": CSVLoader
    }
    if not os.path.exists(docs_dir):
        logger.warning(f"Папка с документами не существует: {docs_dir}")
        os.makedirs(docs_dir, exist_ok=True)
        return documents
    for file in os.listdir(docs_dir):
        file_path = os.path.join(docs_dir, file)
        file_ext = os.path.splitext(file)[1].lower()
        logger.info(f"Обработка файла: {file_path}")
        if file_ext in loaders:
            try:
                text, metadata = loaders[file_ext](file_path)
                logger.info(f"Длина извлечённого текста для {file}: {len(text)}")
                if text.strip():
                    documents.append(Document(page_content=text, metadata=metadata))
                    if file_ext == ".pdf":
                        logger.info(f"Загружен {file} с {len(metadata.get('page_numbers', '').split(','))} страницами")
                    else:
                        logger.info(f"Загружен {file} с содержимым")
                else:
                    logger.warning(f"Нет пригодного текста в {file}")
            except Exception as e:
                logger.error(f"Ошибка загрузки {file}: {e}")
    logger.info(f"Загружено {len(documents)} документов")
    return documents


docs_dir = "/app/data"  
loaded_documents = load_documents(docs_dir)
if loaded_documents:
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=200)
    for doc in loaded_documents:
        chunks = text_splitter.split_text(doc.page_content)
        chunk_embeddings = embeddings.embed_documents(chunks)
        questions = []
        valid_chunk_indices = []  
        
        for i, chunk in enumerate(chunks):
            existing_chunk_ids = [id for id in collection.get()['ids'] if f"chunk_{doc.metadata['source']}_{i}" in id]
            
            if existing_chunk_ids:
                existing_metadata = collection.get(ids=existing_chunk_ids, include=["metadatas"])["metadatas"][0]
                if "questions" in existing_metadata and existing_metadata["questions"] and existing_metadata["questions"] != "[]":
                    questions.append(existing_metadata["questions"])
                    valid_chunk_indices.append(i)
                    logger.info(f"Пропущен чанк {i} с существующими вопросами: {existing_metadata['questions']}")
                continue  
            
            new_questions = generate_questions(chunk)
            if not new_questions or new_questions == []:  
                logger.error(f"Пустые вопросы сгенерированы для чанка {i} после 20 попыток. Остановка программы.")
                # Сохраняем данные только для обработанных чанков (до текущего)
                if questions and valid_chunk_indices and chunks and chunk_embeddings:
                    filtered_chunks = [chunks[j] for j in valid_chunk_indices]
                    filtered_embeddings = [chunk_embeddings[j] for j in valid_chunk_indices]
                    collection.add(
                        documents=filtered_chunks,
                        embeddings=filtered_embeddings,
                        metadatas=[{"questions": "; ".join(q) if isinstance(q, list) else q, "filename": doc.metadata["source"], "type": "chunk", **doc.metadata} for q in questions],
                        ids=[f"chunk_{doc.metadata['source']}_{j}" for j in valid_chunk_indices]
                    )
                    question_texts = ["; ".join(q) if isinstance(q, list) else q for q in questions]
                    collection.add(
                        documents=question_texts,
                        embeddings=embeddings.embed_documents(question_texts) if question_texts else [[]],
                        metadatas=[{"questions": "; ".join(q) if isinstance(q, list) else q, "filename": doc.metadata["source"], "type": "question", "chunk_id": f"chunk_{doc.metadata['source']}_{j}"} for j, q in enumerate(questions)],
                        ids=[f"question_{doc.metadata['source']}_{j}" for j in valid_chunk_indices]
                    )
                    logger.info(f"Сохранены данные для {len(valid_chunk_indices)} чанков до чанка {i}.")
                raise ValueError(f"Пустые вопросы для чанка {i} в файле {doc.metadata['source']}. Обработка прервана. Чанк не сохранён.")
            questions.append(new_questions)
            valid_chunk_indices.append(i)
            logger.info(f"Сгенерированы вопросы для чанка {i}: {new_questions}")

        # Сохранение данных после цикла, если все чанки обработаны
        if chunks and chunk_embeddings and questions:
            filtered_chunks = [chunks[i] for i in valid_chunk_indices]
            filtered_embeddings = [chunk_embeddings[i] for i in valid_chunk_indices]
            collection.add(
                documents=filtered_chunks,
                embeddings=filtered_embeddings,
                metadatas=[{"questions": "; ".join(q) if isinstance(q, list) else q, "filename": doc.metadata["source"], "type": "chunk", **doc.metadata} for q in questions],
                ids=[f"chunk_{doc.metadata['source']}_{i}" for i in valid_chunk_indices]
            )
            question_texts = ["; ".join(q) if isinstance(q, list) else q for q in questions]
            collection.add(
                documents=question_texts,
                embeddings=embeddings.embed_documents(question_texts) if question_texts else [[]],
                metadatas=[{"questions": "; ".join(q) if isinstance(q, list) else q, "filename": doc.metadata["source"], "type": "question", "chunk_id": f"chunk_{doc.metadata['source']}_{i}"} for i, q in enumerate(questions)],
                ids=[f"question_{doc.metadata['source']}_{i}" for i in valid_chunk_indices]
            )
    logger.info(f"Обработано {len(loaded_documents)} документов в коллекцию Chroma")


async def handle_query(question: str):
    try:
        logger.info(f"Получен вопрос: {question}")
        if not hasattr(embeddings, 'embed_query'):
            logger.error("Модель эмбеддингов не работает")
            return {"question": question, "answer": "Ошибка: модель эмбеддингов не работает", "context": "", "used_chunks": [], "used_questions": []}
        
        question_embedding = embeddings.embed_query(question)
        logger.info("Сгенерировано вложение вопроса")

        question_results = collection.query(
            query_embeddings=[question_embedding],
            where={"type": "question"},
            n_results=5,  # Оставляем 5
            include=["documents", "metadatas", "distances", "data"]
        )
        chunk_results = collection.query(
            query_embeddings=[question_embedding],
            where={"type": "chunk"},
            n_results=5,  # Оставляем 5
            include=["documents", "metadatas", "distances"]
        )

        relevant_chunks = []
        
        # Фильтр по вопросам из question_results
        if question_results["metadatas"]:
            question_words = set(re.findall(r'\w+', question.lower()))
            for i, meta in enumerate(question_results["metadatas"][0]):
                if "chunk_id" in meta:
                    chunk_id = meta["chunk_id"]
                    distance = question_results["distances"][0][i]
                    if distance < 0.8:
                        chunk_data = collection.get(ids=[chunk_id], include=["documents", "metadatas"])
                        if chunk_data["documents"]:
                            # Проверка релевантности по вопросам
                            questions = meta.get("questions", "").lower().split("; ")
                            if any(question_words & set(re.findall(r'\w+', q)) for q in questions):
                                relevant_chunks.append({
                                    "text": chunk_data["documents"][0],
                                    "source": chunk_data["metadatas"][0].get("filename", ""),
                                    "score": 1 - distance + 0.1  # Добавить бонус за релевантность вопросов
                                })
                            else:
                                relevant_chunks.append({
                                    "text": chunk_data["documents"][0],
                                    "source": chunk_data["metadatas"][0].get("filename", ""),
                                    "score": 1 - distance
                                })

        # Фильтр по прямым чанкам
        if chunk_results["documents"]:
            question_words = set(re.findall(r'\w+', question.lower()))
            for i, doc in enumerate(chunk_results["documents"][0]):
                distance = chunk_results["distances"][0][i]
                if distance < 0.8:
                    meta = chunk_results["metadatas"][0][i]
                    questions = meta.get("questions", "").lower().split("; ")
                    if any(question_words & set(re.findall(r'\w+', q)) for q in questions):
                        relevant_chunks.append({
                            "text": doc,
                            "source": meta.get("filename", ""),
                            "score": 1 - distance + 0.1  # Бонус за релевантность
                        })
                    else:
                        relevant_chunks.append({
                            "text": doc,
                            "source": meta.get("filename", ""),
                            "score": 1 - distance
                        })

        if not relevant_chunks and chunk_results["documents"]:
            for i, doc in enumerate(chunk_results["documents"][0]):
                relevant_chunks.append({
                    "text": doc,
                    "source": chunk_results["metadatas"][0][i].get("filename", ""),
                    "score": 1 - chunk_results["distances"][0][i]
                })

        unique_chunks = {}
        for chunk in relevant_chunks:
            key = chunk["text"][:500]
            if key not in unique_chunks or chunk["score"] > unique_chunks[key]["score"]:
                unique_chunks[key] = chunk

        sorted_chunks = sorted(unique_chunks.values(), key=lambda x: x["score"], reverse=True)
        context = " ".join([chunk["text"] for chunk in sorted_chunks[:5]])

        if not context:
            logger.info("Релевантные чанки не найдены")
            return {
                "question": question,
                "answer": "Нет данных",
                "context": "",
                "used_chunks": [],
                "used_questions": []
            }

        API_KEY = os.getenv("GOOGLE_API_KEY")
        if not API_KEY:
            return {"answer": "Ошибка: API ключ не настроен"}
            
        BASE_URL = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.0-flash:generateContent?key={API_KEY}"
        
        prompt = f"""
        Ты - юридический ассистент Казахстана. Отвечай ТОЛЬКО на основе предоставленного контекста. 
        Если ответа нет в контексте или контекст не относится к вопросу, верни только строку "Нет данных" и ничего больше. 
        Не генерируй дополнительные вопросы, предположения или структуры (например, JSON). 
        Не добавляй информацию, не упомянутую в контексте.

        Контекст:
        {context[:25000]}

        Вопрос: {question}

        Ответ (только факты из контекста, без дополнений):
        """
        
        payload = {
            "contents": [{"parts": [{"text": prompt}]}],
            "generationConfig": {
                "temperature": 0.1,
                "maxOutputTokens": 1000
            }
        }

        response = requests.post(BASE_URL, json=payload, timeout=60)
        response.raise_for_status()
        response_json = response.json()
        
        if "candidates" in response_json and response_json["candidates"]:
            answer = response_json["candidates"][0]["content"]["parts"][0]["text"]
            if answer.strip().startswith("```") or "follow_ups" in answer.lower() or "{" in answer or "[" in answer:
                logger.warning("Модель вернула некорректный формат, заменяем на 'Нет данных'")
                answer = "Нет данных"
        else:
            answer = "Ошибка: не удалось получить ответ от модели"
        
        logger.info(f"Prompt: {prompt[:500]}...")
        logger.info(f"Answer: {answer}")
        logger.info(f"Used chunks: {[chunk['text'][:200] + '...' for chunk in sorted_chunks[:5]]}")
        logger.info(f"Used questions: {question_results['documents'][0] if question_results['documents'] else 'None'}")

        return {
            "question": question,
            "answer": answer,
            "context": context,
            "used_chunks": [chunk["text"] for chunk in sorted_chunks[:5]],
            "used_questions": question_results["documents"][0] if question_results["documents"] else []
        }

    except Exception as e:
        logger.error(f"Ошибка обработки запроса: {e}", exc_info=True)
        return {
            "question": question,
            "answer": f"Ошибка: {str(e)}",
            "context": "",
            "used_chunks": [],
            "used_questions": []
        }
    


@app.post("/query")
async def query(request: QueryRequest):
    return await handle_query(request.question)

@app.post("/v1/chat/completions")
async def openai_compatible(body: dict):
    try:
        user_message = body.get("messages", [])[-1].get("content")
        result = await handle_query(user_message)
        return {
            "id": "chatcmpl-fake-id",
            "object": "chat.completion",
            "created": int(time.time()),
            "model": "gemini-2.0-flash",
            "choices": [{"index": 0, "message": {"role": "assistant", "content": result["answer"]}, "finish_reason": "stop"}]
        }
    except Exception as e:
        logger.error(f"Ошибка в openai_compatible: {e}")
        raise HTTPException(status_code=500, detail=f"Ошибка совместимости: {e}")

@app.get("/v1/models")
async def get_models():
    return {"object": "list", "data": [{"id": "gemini-2.0-flash", "object": "model", "created": int(time.time()), "owned_by": "xAI"}]}
